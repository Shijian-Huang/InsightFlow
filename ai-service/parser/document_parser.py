import re
from pathlib import Path

from parser.pdf_parser import extract_pdf_title, parse_pdf_pages


SUPPORTED_EXTENSIONS = {".pdf", ".md", ".markdown", ".txt", ".docx"}


def parse_document_pages(path: str) -> list[dict]:
    ext = Path(path).suffix.lower()

    if ext == ".pdf":
        return parse_pdf_pages(path)

    if ext in {".md", ".markdown"}:
        text = Path(path).read_text(encoding="utf-8")
        cleaned = _strip_markdown(text)
        return [{"page": 1, "raw_text": text, "text": cleaned}] if cleaned.strip() else []

    if ext == ".txt":
        text = Path(path).read_text(encoding="utf-8")
        return [{"page": 1, "raw_text": text, "text": text}] if text.strip() else []

    if ext == ".docx":
        text = _extract_docx_text(path)
        return [{"page": 1, "raw_text": text, "text": text}] if text.strip() else []

    raise ValueError(f"Unsupported file format: {ext}")


def extract_document_title(path: str, pages: list[dict] | None = None) -> str:
    ext = Path(path).suffix.lower()

    if ext == ".pdf":
        return extract_pdf_title(path, pages)

    if ext in {".md", ".markdown"}:
        text = Path(path).read_text(encoding="utf-8") if pages is None else (pages[0]["raw_text"] if pages else "")
        for line in text.splitlines():
            stripped = line.strip()
            if stripped.startswith("# ") and not stripped.startswith("##"):
                return stripped.lstrip("# ").strip()
        for line in text.splitlines():
            if line.strip():
                return line.strip()
        return Path(path).stem

    if ext == ".docx":
        return _extract_docx_title(path, pages)

    if ext == ".txt":
        text = pages[0]["text"] if pages else Path(path).read_text(encoding="utf-8")
        for line in text.splitlines():
            if line.strip():
                return line.strip()
        return Path(path).stem

    return Path(path).stem


def _strip_markdown(text: str) -> str:
    result = text
    result = re.sub(r"```[\s\S]*?```", "", result)
    result = re.sub(r"`[^`\n]+`", "", result)
    result = re.sub(r"^#{1,6}\s+", "", result, flags=re.MULTILINE)
    result = re.sub(r"\*\*(.+?)\*\*", r"\1", result)
    result = re.sub(r"__(.+?)__", r"\1", result)
    result = re.sub(r"\*(.+?)\*", r"\1", result)
    result = re.sub(r"_(.+?)_", r"\1", result)
    result = re.sub(r"!?\[([^\]]*)\]\([^)]+\)", r"\1", result)
    result = re.sub(r"^>\s?", "", result, flags=re.MULTILINE)
    result = re.sub(r"^[-*+]\s+", "", result, flags=re.MULTILINE)
    result = re.sub(r"^\d+\.\s+", "", result, flags=re.MULTILINE)
    result = re.sub(r"^---+$", "", result, flags=re.MULTILINE)
    return result


def _extract_docx_text(path: str) -> str:
    from docx import Document

    doc = Document(path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def _extract_docx_title(path: str, pages: list[dict] | None = None) -> str:
    from docx import Document

    doc = Document(path)

    metadata_title = (doc.core_properties.title or "").strip()
    if metadata_title:
        return metadata_title

    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name and paragraph.style.name.startswith("Heading"):
            text = paragraph.text.strip()
            if text:
                return text

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            return text

    return Path(path).stem
